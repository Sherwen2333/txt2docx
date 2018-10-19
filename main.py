from docx import Document
from docx.shared import Pt
document = Document()
def read():
    document.add_heading('Word_1', 0)
    for i in open('word_1.txt','r').readlines():
        i=i.replace('\n','')
        p = document.add_paragraph()
        p.space_after = Pt(5)
        p.space_before = Pt(5)
        if(i.__contains__(':')) and len(i)<20:
            z=p.add_run(i.replace(':', '').capitalize())
            z.bold=True
            z.italic=True
            z.font.size = Pt(30)
        # elif(i=='----------------------------------\n'):
        #     continue
        else:
            p.add_run(i)
    document.save('word_1.docx')

    return
if __name__ == '__main__':
    read()

